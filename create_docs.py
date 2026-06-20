"""
Subterranean Orkestra - .docx Dosya Oluşturucu
İki ana dokümanı Word formatında oluşturur.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def setup_styles(doc):
    """Doküman stillerini ayarla"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Başlık stilleri
    for i in range(1, 4):
        heading = doc.styles[f'Heading {i}']
        heading.font.color.rgb = RGBColor(0, 51, 102)

def add_table(doc, headers, rows):
    """Tablo ekle"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Başlıklar
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Satırlar
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    doc.add_paragraph()

def create_master_instructions():
    """1. Doküman: Master Geliştirme Talimatı"""
    doc = Document()
    setup_styles(doc)
    
    # Kapak
    title = doc.add_heading('AI Asistanı İçin Master Geliştirme Talimatı', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('"Subterranean Orkestra" — Çoklu-Ajan Orkestrasyon Platformu')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph(f'Oluşturulma: 20 Haziran 2026')
    doc.add_page_break()
    
    # BÖLÜM 1
    doc.add_heading('BÖLÜM 1: MISYON, VİZYON VE TEMEL PRENSİPLER', 1)
    
    doc.add_heading('1.1 Projenin Özü', 2)
    doc.add_paragraph(
        'Görevin; "Subterranean Orchestration" felsefesini merkeze alan, kurumsal seviyede, '
        'görsel arayüzlü, çoklu AI ajanını orkestre eden bir masaüstü/web uygulaması geliştirmektir.'
    )
    
    p = doc.add_paragraph()
    p.add_run('Temel Paradigma: ').bold = True
    p.add_run('"Kalıcı yapı ağırlıklara (weights) aittir, geçici durum prompt\'a."')
    
    bullets = [
        'Geleneksel "yüzey orkestrasyonu" (LangGraph/CrewAI ile her adımda LLM\'e talimat gönderme) varsayılan yol OLMAMALIDIR.',
        'Kullanıcı bir iş akışını "Derle" butonuyla küçük bir modele (Qwen-2.5-3B/7B, Llama-3-8B) ince ayar yaparak orkestratörsüz çalıştırılabilir hale getirebilmelidir.',
        'Hedef: Yüzey orkestrasyonuna göre 100-462x daha ucuz, frontier model kalitesinin %87-98\'ine ulaşan derlenmiş ajanlar.'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_heading('1.2 Tasarım Prensipleri', 2)
    principles = [
        ('Maliyet Bilinci', 'Her kod satırında token maliyetini düşün. Gereksiz LLM çağrısı yapma.'),
        ('Modülerlik', 'Her bileşen (model gateway, orchestrator, memory) bağımsız değiştirilebilir olmalı.'),
        ('Gözlemlenebilirlik Varsayılanı', 'İzlenemeyen hiçbir özellik üretilmez.'),
        ('Güvenlik by Design', 'API anahtarları, kullanıcı verileri şifreli; hassas işlemler yerel çalıştırılabilir.'),
        ('Kullanıcı Özerkliği', 'Kullanıcıya her zaman "duraklat, müdahale et, derle, dağıt" yetkisi ver (HITL).')
    ]
    for i, (title, desc) in enumerate(principles, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)
    
    # BÖLÜM 2
    doc.add_heading('BÖLÜM 2: TEKNOLOJİ YIĞINI', 1)
    
    tech_stack = [
        ['Backend API', 'Python 3.11+ / FastAPI', 'Asenkron, OpenTelemetry entegrasyonu'],
        ['Orkestrasyon', 'LangGraph (birincil) + CrewAI', 'Durum bilgisi, hata kurtarma'],
        ['Model Gateway', 'LiteLLM veya özel model_gateway.py', 'Çoklu sağlayıcı tek API'],
        ['İlişkisel DB', 'PostgreSQL 16 + SQLAlchemy 2.0', 'Oturum, ajan meta verisi, RBAC'],
        ['Vektör DB', 'Qdrant (tercih) / Pinecone', 'RAG ve uzun vadeli bellek'],
        ['Mesaj Kuyruğu', 'Redis Streams veya RabbitMQ', 'Olay tabanlı tetikleyiciler'],
        ['GUI', 'Tauri 2.0 + SvelteKit (TypeScript)', 'Electron\'dan %80 daha hafif'],
        ['İletişim', 'REST + WebSocket + SSE', 'Canlı akış, HITL onayları'],
        ['Konteyner', 'Docker + Kubernetes (Helm)', 'Auto-scaling, federe dağıtım'],
        ['Gözlemlenebilirlik', 'OpenTelemetry + Grafana + Jaeger', 'Uçtan uca izleme, FinOps'],
        ['Derleme Altyapısı', 'Unsloth + Hugging Face TRL', 'LoRA ince ayar, 30-50 dk']
    ]
    add_table(doc, ['Katman', 'Seçim', 'Gerekçe'], tech_stack)
    
    # BÖLÜM 3
    doc.add_heading('BÖLÜM 3: MİMARİ VE MODÜLLER', 1)
    
    doc.add_heading('3.1 Çekirdek Modüller (Backend)', 2)
    modules = [
        'core/model_gateway.py — Tüm LLM sağlayıcıları için tek generate()',
        'core/orchestrator.py — LangGraph tabanlı iş akışı yürütücü',
        'core/memory_manager.py — Kısa/uzun vadeli bellek (PG + Qdrant)',
        'core/subterranean_compiler.py — ⭐ KRİTİK: Derleme motoru',
        'core/security.py — JWT + RBAC + şifreleme',
        'agents/base_agent.py — Soyut ajan sınıfı',
        'agents/registry.py — Ajan keşif ve kayıt (A2A protokolü)',
        'workflows/patterns/ — Sequential, Parallel, Supervisor, Feedback, Federated',
        'triggers/ — cron.py, webhook.py, file_watcher.py, queue.py',
        'observability/tracing.py — OpenTelemetry',
        'observability/finops.py — Token, maliyet, bütçe takibi'
    ]
    for m in modules:
        doc.add_paragraph(m, style='List Bullet')
    
    doc.add_heading('3.2 Desteklenen Orkestrasyon Desenleri', 2)
    patterns = [
        ['Supervisor-Worker (Merkezi)', 'Müşteri hizmetleri, bankacılık'],
        ['Sequential Pipeline', 'Araştırma → Taslak → Yazma → Kontrol'],
        ['Parallel Execution', 'Bağımsız alt görevler (ödeme + dolandırıcılık + iade)'],
        ['Feedback Loop', 'Kod yazma + güvenlik denetimi (döngüsel)'],
        ['Federated', 'Çok bulutlu, merkeziyetsiz kurumsal ortamlar']
    ]
    add_table(doc, ['Desen', 'Kullanım Yeri'], patterns)
    
    # BÖLÜM 4
    doc.add_heading('BÖLÜM 4: GELİŞTİRME AŞAMALARI (12 AY)', 1)
    
    phases = [
        ('🟢 Aşama 1: Çekirdek Motor (0-3 Ay)', 
         'Başarı Kriteri: CLI üzerinden agenta run workflow.json ile sıralı bir iş akışını çalıştırıp sonuçları görebilmek.',
         ['model_gateway.py — 5 sağlayıcı desteği', 
          'orchestrator.py — Sequential + Parallel desenleri',
          'memory_manager.py — PG + Qdrant, oturum bazlı bellek',
          'security.py — JWT + admin/user rolleri',
          'cli.py — agenta create|run|list|logs|compile komutları',
          'Birim testleri (%80+ coverage)']),
        ('🟡 Aşama 2: Görsel Orkestrasyon (3-6 Ay)',
         'Başarı Kriteri: Sürükle-bırak ile oluşturulan akışın "Çalıştır" ile başlatılıp canlı izlenebilmesi.',
         ['FlowCanvas.svelte — React Flow benzeri canvas',
          'Özellik paneli (model, prompt, I/O değişkenleri)',
          'Dashboard — Proje, ajan, son çalıştırmalar listesi',
          'RAG modülü — PDF/TXT/MD yükleme, chunking',
          'WebSocket entegrasyonu — Adım adım canlı ilerleme',
          'Etkileşimli hata ayıklayıcı']),
        ('🟠 Aşama 3: Kurumsal Özellikler (6-9 Ay)',
         'Başarı Kriteri: Subterranean derlenmiş bir modelin orkestratörsüz çalışması + 100+ eşzamanlı ajan.',
         ['RBAC — Admin/Manager/Developer/Viewer + SSO/LDAP',
          'Tetikleyici sistemi — Cron, webhook, file watcher, RabbitMQ',
          'HITL — await_approval düğümü, GUI\'de onay/red/düzenle',
          '⭐ subterranean_compiler.py — Tam derleme pipeline',
          'FinOps paneli — Token, maliyet, süre, başarı oranı',
          'OpenTelemetry ile dağıtık izleme']),
        ('🔴 Aşama 4: Ölçekleme ve Dağıtım (9-12 Ay)',
         'Başarı Kriteri: Canlı ortamda 1000+ eşzamanlı iş akışı, 3 iterasyon iyileştirme.',
         ['Dockerfile + Kubernetes Helm chart',
          'Federe mimari — mTLS/VPN ile ajanlar arası güvenli iletişim',
          'CI/CD — GitHub Actions pipeline',
          'PWA mobil arayüz',
          'OpenAPI/Swagger dokümantasyonu'])
    ]
    
    for title, criteria, deliverables in phases:
        doc.add_heading(title, 2)
        p = doc.add_paragraph()
        p.add_run(criteria).italic = True
        for d in deliverables:
            doc.add_paragraph(d, style='List Bullet')
    
    # BÖLÜM 5
    doc.add_heading('BÖLÜM 5: KOD STANDARTLARI', 1)
    standards = [
        'Python: PEP8, ruff linter, black formatter, mypy tip kontrolü',
        'TypeScript/Svelte: ESLint + Prettier, strict mode',
        'Commit: Conventional Commits (feat:, fix:, docs:)',
        'Branch: GitHub Flow (main + feature branch\'ler)',
        'Her yeni özellik için birim + entegrasyon testi',
        'Coverage %80\'in altına düşmez',
        'Her modül için docstring (Google style)',
        'Her API endpoint için OpenAPI schema'
    ]
    for s in standards:
        doc.add_paragraph(s, style='List Bullet')
    
    # BÖLÜM 6
    doc.add_heading('BÖLÜM 6: BAŞARI METRİKLERİ (KPIs)', 1)
    kpis = [
        ['İş akışı başarı oranı', '≥ %95', 'Prometheus counter'],
        ['Ortalama ajan yanıt süresi', '≤ 3 sn', 'OpenTelemetry histogram'],
        ['İş akışı tamamlanma süresi', '≤ 10 sn', 'Uçtan uca trace'],
        ['Subterranean maliyet tasarrufu', '≥ 100x', 'FinOps paneli'],
        ['Derleme süresi', '30-50 dk', 'Compiler progress'],
        ['GUI sezgisellik puanı', '≥ %80', 'Beta anketi'],
        ['Test coverage', '≥ %80', 'pytest-cov / jest']
    ]
    add_table(doc, ['Metrik', 'Hedef', 'Ölçüm Yöntemi'], kpis)
    
    # BÖLÜM 7
    doc.add_heading('BÖLÜM 7: RİSKLER VE TUZAKLAR', 1)
    risks = [
        ('"Ajan Sprawl" Tuzağı', 'Kontrolsüz ajan ekleme. Her ajan için net yetki alanı tanımla.'),
        ('Prompt Enjeksiyonu', 'Tüm kullanıcı girdilerini sanitize et, güven zinciri kur.'),
        ('Maliyet Patlaması', 'Her proje için bütçe limiti zorunlu. Limit aşımı otomatik durdurma.'),
        ('Vendor Lock-in', 'Model gateway soyutlaması ASLA atlanmaz.'),
        ('Bellek Şişmesi', 'Vektör DB\'de TTL ve boyut limitleri tanımla.'),
        ('Derleme Başarısızlığı', 'Kalite doğrulama (sentetik test seti) zorunlu.')
    ]
    for title, desc in risks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'❌ {title}: ').bold = True
        p.add_run(desc)
    
    # Kaydet
    filename = 'Subterranean_Orkestra_Master_Talimatlar.docx'
    doc.save(filename)
    print(f'✓ Oluşturuldu: {filename}')
    return filename

def create_implementation_guide():
    """2. Doküman: Talimat Uygulama Rehberi"""
    doc = Document()
    setup_styles(doc)
    
    # Kapak
    title = doc.add_heading('Subterranean Orkestra — Talimat Uygulama Rehberi', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Her talimatın somut olarak nasıl gerçekleştirileceği')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph(f'Oluşturulma: 20 Haziran 2026')
    doc.add_page_break()
    
    # BÖLÜM 1
    doc.add_heading('BÖLÜM 1: Misyon ve Temel Prensipler — Uygulama', 1)
    
    doc.add_heading('1.1 Subterranean Felsefesini Koda Dökmek', 2)
    doc.add_paragraph(
        'Prensip: "Kalıcı yapı ağırlıklara aittir, geçici durum prompt\'a."'
    )
    
    doc.add_heading('SubterraneanCompiler Sınıfı', 3)
    compiler_code = """class SubterraneanCompiler:
    async def compile(self, workflow_id: str, target_model: str = "qwen2.5-7b"):
        # 1. İş akışını yükle
        workflow = await self.workflow_repo.get(workflow_id)
        
        # 2. Claude/GPT-4 ile sentetik diyalog üret (1000+ örnek)
        synthetic_data = await self.data_generator.generate(
            workflow=workflow,
            num_samples=1500,
            diversity_temperature=0.8
        )
        
        # 3. Unsloth ile LoRA fine-tune
        lora_adapter = await self.fine_tuner.train(
            base_model=target_model,
            dataset=synthetic_data,
            lora_rank=16,
            epochs=3,
            progress_callback=self._emit_progress
        )
        
        # 4. vLLM ile endpoint deploy et
        endpoint = await self.deployer.deploy(
            adapter=lora_adapter,
            base_model=target_model
        )
        
        # 5. Model gateway'e kaydet
        await self.model_gateway.register_compiled_model(
            workflow_id=workflow_id,
            endpoint_url=endpoint.url
        )
        
        return endpoint"""
    
    p = doc.add_paragraph()
    run = p.add_run(compiler_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('1.2 Maliyet Bilinci — Her Satırda Token Hesabı', 2)
    cost_code = """def track_cost(model_name: str):
    def decorator(func):
        @wraps(func)
        async def wrapper(*args, **kwargs):
            start = time.time()
            result = await func(*args, **kwargs)
            
            prompt_tokens = count_tokens(kwargs.get("prompt", ""))
            completion_tokens = count_tokens(result.text)
            cost = calculate_cost(model_name, prompt_tokens, completion_tokens)
            
            COST_COUNTER.labels(
                model=model_name,
                project=kwargs.get("project_id")
            ).inc(cost)
            
            if await is_budget_exceeded(kwargs.get("project_id")):
                raise BudgetExceededError("Proje bütçesi aşıldı")
            
            return result
        return wrapper
    return decorator"""
    
    p = doc.add_paragraph()
    run = p.add_run(cost_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # BÖLÜM 2
    doc.add_heading('BÖLÜM 2: Teknoloji Yığını — Kurulum', 1)
    
    doc.add_heading('2.1 Proje İskeleti Oluşturma', 2)
    setup_code = """# Proje dizini
mkdir -p subterranean-orchestra/{backend,frontend,infra,docs}
cd subterranean-orchestra

# Backend (Python)
cd backend
python -m venv venv
source venv/bin/activate

# Gerekli paketler
pip install fastapi==0.111.0 uvicorn[standard] langgraph==0.2.0 \\
    litellm==1.40.0 sqlalchemy==2.0.30 asyncpg qdrant-client==1.9.0 \\
    redis==5.0.0 pydantic==2.7.0 python-jose[cryptography] \\
    passlib[bcrypt] opentelemetry-api opentelemetry-sdk \\
    prometheus-client unsloth

pip freeze > requirements.txt

# Frontend (Tauri + Svelte)
cd ../frontend
npm create tauri-app@latest . -- --template svelte-ts"""
    
    p = doc.add_paragraph()
    run = p.add_run(setup_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('2.2 Docker Compose', 2)
    docker_code = """version: '3.9'
services:
  postgres:
    image: postgres:16-alpine
    environment:
      POSTGRES_DB: orchestra
      POSTGRES_USER: orchestra
      POSTGRES_PASSWORD: ${DB_PASSWORD}
    ports: ["5432:5432"]

  qdrant:
    image: qdrant/qdrant:v1.9.0
    ports: ["6333:6333"]

  redis:
    image: redis:7-alpine
    ports: ["6379:6379"]

  jaeger:
    image: jaegertracing/all-in-one:1.57
    ports:
      - "16686:16686"
      - "4317:4317"

  prometheus:
    image: prom/prometheus:v2.51.0
    ports: ["9090:9090"]

  grafana:
    image: grafana/grafana:10.4.0
    ports: ["3000:3000"]

  vllm:
    image: vllm/vllm-openai:latest
    runtime: nvidia
    deploy:
      resources:
        reservations:
          devices:
            - driver: nvidia
              count: 1
              capabilities: [gpu]
    ports: ["8000:8000"]"""
    
    p = doc.add_paragraph()
    run = p.add_run(docker_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # BÖLÜM 3
    doc.add_heading('BÖLÜM 3: Mimari ve Modüller — Kodlama', 1)
    
    doc.add_heading('3.1 Model Gateway', 2)
    doc.add_paragraph(
        'Tüm LLM sağlayıcıları için tek bir generate() fonksiyonu. '
        'OpenAI, Anthropic, Google, Azure, Ollama ve derlenmiş modelleri destekler.'
    )
    
    gateway_code = """class ModelGateway:
    PROVIDER_MAP = {
        "openai": "openai",
        "anthropic": "anthropic",
        "google": "gemini",
        "azure": "azure",
        "ollama": "ollama",
        "compiled": "openai"  # vLLM endpoint'leri OpenAI uyumlu
    }
    
    async def generate(self, prompt: str, model: str, 
                       provider: str = "openai",
                       temperature: float = 0.7,
                       stream: bool = False):
        if provider == "compiled":
            base_url = await self._get_compiled_endpoint(model)
            full_model = f"openai/{model}"
            kwargs = {"api_base": base_url}
        else:
            full_model = f"{provider}/{model}"
            kwargs = {}
        
        response = await acompletion(
            model=full_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
            stream=stream,
            **kwargs
        )
        return response.choices[0].message.content"""
    
    p = doc.add_paragraph()
    run = p.add_run(gateway_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('3.2 Orchestrator (LangGraph)', 2)
    orchestrator_code = """class WorkflowState(TypedDict):
    messages: Annotated[list, operator.add]
    current_agent: str
    context: dict
    iteration: int
    status: Literal["running", "awaiting_approval", "completed", "failed"]
    compiled_model: str | None

class Orchestrator:
    def build_sequential_workflow(self, agents: list[dict]) -> StateGraph:
        workflow = StateGraph(WorkflowState)
        
        for agent in agents:
            workflow.add_node(agent["id"], self._make_agent_node(agent))
        
        for i in range(len(agents) - 1):
            workflow.add_edge(agents[i]["id"], agents[i+1]["id"])
        
        workflow.set_entry_point(agents[0]["id"])
        workflow.add_edge(agents[-1]["id"], END)
        
        return workflow.compile()
    
    async def run(self, workflow, initial_state: dict):
        async for event in workflow.astream(initial_state):
            await self._emit_progress(event)
            yield event"""
    
    p = doc.add_paragraph()
    run = p.add_run(orchestrator_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('3.3 Memory Manager', 2)
    memory_code = """class MemoryManager:
    def __init__(self, db: AsyncSession, qdrant: AsyncQdrantClient):
        self.db = db
        self.qdrant = qdrant
        self.embedder = SentenceTransformer("all-MiniLM-L6-v2")
    
    async def save_turn(self, session_id: str, agent_id: str, 
                        input: str, output: str):
        # Kısa vadeli: PostgreSQL
        turn = ConversationTurn(
            session_id=session_id,
            agent_id=agent_id,
            input=input,
            output=output,
            timestamp=datetime.utcnow()
        )
        self.db.add(turn)
        await self.db.commit()
        
        # Uzun vadeli: Vektör DB
        embedding = self.embedder.encode(f"{input} | {output}")
        await self.qdrant.upsert(
            collection_name="long_term_memory",
            points=[PointStruct(
                id=str(uuid.uuid4()),
                vector=embedding.tolist(),
                payload={
                    "session_id": session_id,
                    "agent_id": agent_id,
                    "text": f"{input}\\n→ {output}",
                    "timestamp": datetime.utcnow().isoformat()
                }
            )]
        )
    
    async def semantic_search(self, query: str, top_k: int = 5):
        query_vec = self.embedder.encode(query).tolist()
        results = await self.qdrant.search(
            collection_name="long_term_memory",
            query_vector=query_vec,
            limit=top_k
        )
        return [r.payload for r in results]"""
    
    p = doc.add_paragraph()
    run = p.add_run(memory_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # BÖLÜM 4
    doc.add_heading('BÖLÜM 4: Geliştirme Aşamaları — Adım Adım', 1)
    
    doc.add_heading('Aşama 1: Çekirdek Motor (Hafta 1-12)', 2)
    phase1 = [
        'Hafta 1-2: Proje iskeleti + Model Gateway',
        'Hafta 3-4: Orchestrator + Memory',
        'Hafta 5-8: CLI + Testler',
        'Hafta 9-12: Entegrasyon testleri + İlk KPI raporu'
    ]
    for item in phase1:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Aşama 2: Görsel Orkestrasyon (Ay 3-6)', 2)
    doc.add_paragraph(
        'Svelte + @xyflow/svelte ile sürükle-bırak canvas. '
        'WebSocket ile canlı ilerleme takibi.'
    )
    
    doc.add_heading('Aşama 3: Kurumsal Özellikler (Ay 6-9)', 2)
    phase3 = [
        'RBAC (Admin/Manager/Developer/Viewer)',
        'HITL (await_approval düğümü)',
        'Subterranean Compiler (tam pipeline)',
        'FinOps paneli',
        'OpenTelemetry dağıtık izleme'
    ]
    for item in phase3:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Aşama 4: Ölçekleme (Ay 9-12)', 2)
    phase4 = [
        'Kubernetes Helm chart',
        'Federe mimari (mTLS/VPN)',
        'CI/CD pipeline',
        'PWA mobil arayüz',
        'OpenAPI/Swagger dokümantasyon'
    ]
    for item in phase4:
        doc.add_paragraph(item, style='List Bullet')
    
    # BÖLÜM 5
    doc.add_heading('BÖLÜM 5: Kod Standartları', 1)
    
    doc.add_heading('5.1 Pre-commit Hooks', 2)
    precommit_code = """repos:
  - repo: https://github.com/astral-sh/ruff-pre-commit
    rev: v0.4.0
    hooks:
      - id: ruff
        args: [--fix]
      - id: ruff-format
  
  - repo: https://github.com/pre-commit/mirrors-mypy
    rev: v1.10.0
    hooks:
      - id: mypy"""
    
    p = doc.add_paragraph()
    run = p.add_run(precommit_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('5.2 Hata Yönetimi Şablonu', 2)
    error_code = """class OrchestraBaseError(Exception):
    def __init__(self, message: str, trace_id: str = None):
        self.message = message
        self.trace_id = trace_id or str(uuid.uuid4())
        super().__init__(self.message)

class ModelRateLimitError(OrchestraBaseError): pass
class BudgetExceededError(OrchestraBaseError): pass
class CompilationQualityError(OrchestraBaseError): pass

try:
    result = await agent.execute(task)
except ModelRateLimitError as e:
    logger.warning(f"Rate limit, fallback'a geçiliyor: {e.trace_id}")
    result = await fallback_agent.execute(task)
except BudgetExceededError as e:
    await alert_ops_channel(f"Bütçe aşıldı: {e.trace_id}")
    raise
except Exception as e:
    logger.exception(f"Beklenmeyen hata: {e}")
    await alert_ops_channel(e)
    raise OrchestraBaseError(str(e)) from e"""
    
    p = doc.add_paragraph()
    run = p.add_run(error_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # BÖLÜM 6
    doc.add_heading('BÖLÜM 6: Başarı Metrikleri', 1)
    
    metrics_code = """from prometheus_client import Counter, Histogram, Gauge

WORKFLOW_STATUS = Counter(
    "orchestra_workflow_status_total",
    "İş akışı durumları",
    ["workflow_id", "status"]
)

AGENT_LATENCY = Histogram(
    "orchestra_agent_latency_seconds",
    "Ajan yanıt süresi",
    ["agent_id", "model"],
    buckets=[0.5, 1, 2, 3, 5, 10, 30, 60]
)

COST_USD = Counter(
    "orchestra_cost_usd_total",
    "Toplam maliyet (USD)",
    ["project_id", "model"]
)

COMPILATION_DURATION = Histogram(
    "orchestra_compilation_seconds",
    "Derleme süresi",
    buckets=[600, 1200, 1800, 2400, 3000]
)"""
    
    p = doc.add_paragraph()
    run = p.add_run(metrics_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # BÖLÜM 7
    doc.add_heading('BÖLÜM 7: Risk Yönetimi', 1)
    
    doc.add_heading('Risk 1: Maliyet Patlaması', 2)
    budget_code = """class BudgetGuard:
    async def check(self, project_id: str, estimated_cost: float):
        current = await self.get_current_spend(project_id)
        limit = await self.get_budget_limit(project_id)
        
        if current + estimated_cost > limit:
            await self.pause_all_workflows(project_id)
            await notify_admin(f"Proje {project_id} bütçeyi aştı!")
            raise BudgetExceededError(f"{current}/{limit} USD")"""
    
    p = doc.add_paragraph()
    run = p.add_run(budget_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('Risk 2: Prompt Enjeksiyonu', 2)
    sanitize_code = """def sanitize_input(user_input: str) -> str:
    dangerous_patterns = [
        r"ignore previous instructions",
        r"you are now",
        r"system:",
        r"<\\|im_start\\|>"
    ]
    for pattern in dangerous_patterns:
        user_input = re.sub(pattern, "[REDACTED]", user_input, flags=re.IGNORECASE)
    return user_input"""
    
    p = doc.add_paragraph()
    run = p.add_run(sanitize_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('Risk 3: Ajan Sprawl', 2)
    registry_code = """class AgentRegistry:
    async def register(self, agent: Agent):
        if not agent.domain or len(agent.domain) < 20:
            raise AgentValidationError("Ajanın net bir yetki alanı olmalı")
        
        existing = await self.count_by_domain(agent.domain)
        if existing >= 5:
            raise AgentValidationError(
                f"'{agent.domain}' alanında zaten {existing} ajan var."
            )
        
        if not await self.run_benchmark_tests(agent):
            raise AgentValidationError("Ajan benchmark testlerini geçemedi")"""
    
    p = doc.add_paragraph()
    run = p.add_run(registry_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # BÖLÜM 8
    doc.add_heading('BÖLÜM 8: İlk 7 Gün Planı', 1)
    
    days = [
        ('Gün 1', 'Proje iskeleti — Tüm klasör yapısını oluştur'),
        ('Gün 2-3', 'Model Gateway + İlk test (pytest)'),
        ('Gün 4-5', 'Orchestrator + Sequential Workflow'),
        ('Gün 6-7', 'İlk KPI raporu oluştur')
    ]
    for day, task in days:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{day}: ').bold = True
        p.add_run(task)
    
    doc.add_heading('İlk hafta teslimatları:', 2)
    deliverables = [
        '✓ model_gateway.py (5 sağlayıcı)',
        '✓ orchestrator.py (Sequential + Parallel)',
        '✓ memory_manager.py (PG + Qdrant)',
        '✓ cli.py (create, run, list, logs)',
        '✓ Docker Compose (tüm altyapı)',
        '✓ İlk KPI raporu'
    ]
    for d in deliverables:
        doc.add_paragraph(d, style='List Bullet')
    
    # Kaydet
    filename = 'Subterranean_Orkestra_Uygulama_Rehberi.docx'
    doc.save(filename)
    print(f'✓ Oluşturuldu: {filename}')
    return filename

if __name__ == "__main__":
    print('📄 Subterranean Orkestra - .docx Dosya Oluşturucu')
    print('=' * 60)
    
    file1 = create_master_instructions()
    file2 = create_implementation_guide()
    
    print('=' * 60)
    print(f'✅ Toplam 2 dosya oluşturuldu:')
    print(f'   1. {file1}')
    print(f'   2. {file2}')
    print('\n💡 Dosyaları Microsoft Word, LibreOffice veya Google Docs ile açabilirsiniz.')